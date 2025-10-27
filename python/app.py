import os
import pandas as pd
from sentence_transformers import SentenceTransformer
from sklearn.preprocessing import normalize
import numpy as np
import time
import faiss
from llama_cpp import Llama
import requests
import json
from dotenv import load_dotenv
import spacy
import nltk
from fastapi import FastAPI, Request, Form , HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import random
import re
from textblob import TextBlob
from spellchecker import SpellChecker 
import threading
import queue  # Added for background task queue
from urllib.parse import quote
import hashlib
from docx import Document as DocxDocument
import io

# ==================== NEW: Background Embedding System ====================
_EMBEDDING_QUEUE = queue.Queue()
_BACKGROUND_THREAD = None
_STOP_BACKGROUND_THREAD = False

class EmbeddingStatus:
    def __init__(self):
        self.lock = threading.Lock()
        self.pending_files = set()
        self.processing_file = None
        self.completed_files = set()
        self.failed_files = {}
        
    def add_pending(self, file_path):
        with self.lock:
            self.pending_files.add(file_path)
    
    def start_processing(self, file_path):
        with self.lock:
            self.pending_files.discard(file_path)
            self.processing_file = file_path
    
    def mark_completed(self, file_path):
        with self.lock:
            self.processing_file = None
            self.completed_files.add(file_path)
    
    def mark_failed(self, file_path, error):
        with self.lock:
            self.processing_file = None
            self.failed_files[file_path] = str(error)
    
    def get_status(self):
        with self.lock:
            return {
                "pending": list(self.pending_files),
                "processing": self.processing_file,
                "completed": list(self.completed_files),
                "failed": dict(self.failed_files)
            }

embedding_status = EmbeddingStatus()

def background_embedding_worker():
    """Background thread that continuously processes embedding requests"""
    global _STOP_BACKGROUND_THREAD
    print("🚀 Background embedding worker started")
    
    while not _STOP_BACKGROUND_THREAD:
        try:
            try:
                file_path = _EMBEDDING_QUEUE.get(timeout=1.0)
            except queue.Empty:
                continue
            
            if file_path is None:  # Poison pill to stop thread
                break
            
            print(f"📝 Processing embeddings for: {os.path.basename(file_path)}")
            embedding_status.start_processing(file_path)
            
            success = _create_embeddings_for_file(file_path)
            
            if success:
                embedding_status.mark_completed(file_path)
                print(f"✅ Completed embeddings for: {os.path.basename(file_path)}")
            else:
                embedding_status.mark_failed(file_path, "Embedding creation failed")
                print(f"❌ Failed embeddings for: {os.path.basename(file_path)}")
            
            _EMBEDDING_QUEUE.task_done()
            
        except Exception as e:
            print(f"❌ Error in background worker: {e}")
            import traceback
            traceback.print_exc()
    
    print("🛑 Background embedding worker stopped")

def _create_embeddings_for_file(file_path):
    """Create embeddings for a single file"""
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return False
        
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        embeddings_dir = os.path.join(path, "embeddings")
        os.makedirs(embeddings_dir, exist_ok=True)
        
        emb_path = os.path.join(embeddings_dir, f"{base_name}_faiss.index")
        chunk_path = os.path.join(embeddings_dir, f"{base_name}_chunks.npy")
        
        # Check if already exists
        if os.path.exists(emb_path) and os.path.exists(chunk_path):
            print(f"Embeddings already exist for {base_name}, skipping")
            return True
        
        # Load and process CSV
        chunks_from_file, _ = multiple_csv([file_path])
        
        if not chunks_from_file:
            print(f"No valid chunks found in {file_path}")
            return False
        
        # Create embedder
        embedder = create_embedding_model()
        
        # Generate embeddings
        print(f"Encoding {len(chunks_from_file)} chunks for {base_name}...")
        embeddings_np = embedder.encode(chunks_from_file, show_progress_bar=True)
        embeddings_normalized = normalize(embeddings_np.astype(np.float32), axis=1, norm='l2')
        
        # Create and save index
        dim = embeddings_normalized.shape[1]
        index = faiss.IndexFlatIP(dim)
        index.add(embeddings_normalized)
        
        faiss.write_index(index, emb_path)
        np.save(chunk_path, np.array(chunks_from_file, dtype=object))
        
        print(f"✅ Saved embeddings for {base_name}")
        return True
        
    except Exception as e:
        print(f"Error creating embeddings for {file_path}: {e}")
        import traceback
        traceback.print_exc()
        return False

def start_background_worker():
    """Start the background embedding worker thread"""
    global _BACKGROUND_THREAD, _STOP_BACKGROUND_THREAD
    
    if _BACKGROUND_THREAD is not None and _BACKGROUND_THREAD.is_alive():
        print("Background worker already running")
        return
    
    _STOP_BACKGROUND_THREAD = False
    _BACKGROUND_THREAD = threading.Thread(target=background_embedding_worker, daemon=True)
    _BACKGROUND_THREAD.start()

def stop_background_worker():
    """Stop the background embedding worker thread"""
    global _BACKGROUND_THREAD, _STOP_BACKGROUND_THREAD
    
    if _BACKGROUND_THREAD is None or not _BACKGROUND_THREAD.is_alive():
        print("Background worker not running")
        return
    
    print("Stopping background worker...")
    _STOP_BACKGROUND_THREAD = True
    _EMBEDDING_QUEUE.put(None)  # Poison pill
    _BACKGROUND_THREAD.join(timeout=5.0)
    print("Background worker stopped")

def queue_file_for_embedding(file_path):
    """Add a file to the embedding queue for background processing"""
    if not os.path.exists(file_path):
        print(f"Cannot queue non-existent file: {file_path}")
        return False
    
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    embeddings_dir = os.path.join(path, "embeddings")
    emb_path = os.path.join(embeddings_dir, f"{base_name}_faiss.index")
    chunk_path = os.path.join(embeddings_dir, f"{base_name}_chunks.npy")
    
    # Check if embeddings already exist
    if os.path.exists(emb_path) and os.path.exists(chunk_path):
        print(f"Embeddings already exist for {base_name}, skipping queue")
        return True
    
    # Add to queue
    embedding_status.add_pending(file_path)
    _EMBEDDING_QUEUE.put(file_path)
    print(f"📥 Queued for embedding: {os.path.basename(file_path)}")
    return True

def create_faiss_index_fast(embedder, file_names):
    """
    FAST version: Only loads existing embeddings, doesn't create new ones
    New files are handled by background worker
    """
    print("Creating FAISS index (fast mode - loading existing only)...")
    
    all_chunks_global = []
    chunk_to_file_map_global = {}
    embeddings_dir = os.path.join(path, "embeddings")
    embedding_dim = embedder.get_sentence_embedding_dimension()
    main_index = faiss.IndexFlatIP(embedding_dim)
    
    files_needing_embeddings = []
    
    for file_path in file_names:
        if not os.path.exists(file_path):
            continue
        
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        emb_path = os.path.join(embeddings_dir, f"{base_name}_faiss.index")
        chunk_path = os.path.join(embeddings_dir, f"{base_name}_chunks.npy")
        
        # Try to load existing embeddings
        if os.path.exists(emb_path) and os.path.exists(chunk_path):
            try:
                individual_index = faiss.read_index(emb_path)
                chunks_loaded = np.load(chunk_path, allow_pickle=True).tolist()
                
                if individual_index.ntotal > 0 and len(chunks_loaded) == individual_index.ntotal:
                    embeddings = np.vstack([individual_index.reconstruct(i) 
                                           for i in range(individual_index.ntotal)]).astype(np.float32)
                    
                    all_chunks_global.extend(chunks_loaded)
                    for chunk in chunks_loaded:
                        chunk_to_file_map_global[chunk] = os.path.basename(file_path)
                    main_index.add(embeddings)
                    
                    print(f"✅ Loaded {base_name}: {individual_index.ntotal} vectors")
                    continue
            except Exception as e:
                print(f"Error loading {base_name}: {e}")
        
        # File needs embeddings - queue it
        files_needing_embeddings.append(file_path)
        print(f"⏳ {base_name}: Embeddings pending")
    
    # Queue files that need embeddings
    for file_path in files_needing_embeddings:
        queue_file_for_embedding(file_path)
    
    return embedder, all_chunks_global, main_index, chunk_to_file_map_global

# ==================== END NEW CODE ====================

# --- Original imports and globals ---
_EMBEDDING_UPDATE_IN_PROGRESS = False
_EMBEDDING_UPDATE_LOCK = threading.Lock()

# Download required NLTK data
nltk.download('punkt', quiet=True)
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')
os.system("python -m textblob.download_corpora")
load_dotenv()
path = ''
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize spaCy
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("spaCy English model 'en_core_web_sm' not found. Downloading...")
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# ==================== ALL YOUR ORIGINAL FUNCTIONS (UNCHANGED) ====================

def entity_overlap(text1, text2):
    """Calculate entity overlap between two texts"""
    doc1 = nlp(text1)
    doc2 = nlp(text2)
    entities1 = set([ent.text.lower() for ent in doc1.ents])
    entities2 = set([ent.text.lower() for ent in doc2.ents])
    if not entities1 or not entities2:
        return 0
    overlap = len(entities1.intersection(entities2))
    total = len(entities1.union(entities2))
    return overlap / total if total > 0 else 0

l = ["I'm sorry, I don't have that information.",
     "I'm not sure about that.?" , 
     "I don't know the answer to that question. Is there something else I can help with?",
     "That information is not available to me.",
     "I'm unable to answer that question right now"]

def extract_questions(text):
    """Enhanced question extraction that preserves proper names and titles"""
    text = text.strip()
    greeting = extract_greeting(text)
    if greeting:
        text = text.replace(greeting, '').strip()
    if '?' in text:
        parts = [p.strip() for p in text.split('?') if p.strip()]
        if len(parts) > 1:
            return [f"{p}?" for p in parts[:-1]] + [parts[-1]]
    doc = nlp(text)
    keep_together = []
    current_span = []
    for token in doc:
        if token.ent_type_ or token.pos_ in ('PROPN', 'NOUN'):
            current_span.append(token.text)
        else:
            if current_span:
                keep_together.append(' '.join(current_span))
                current_span = []
    if current_span:
        keep_together.append(' '.join(current_span))
    standardized = text.lower()
    conjunctions = [
        (r'\band\b', 'AND'),
        (r'\btell me about\b', 'TELLME'),
        (r'\bwhat about\b', 'WHATABOUT'),
        (r'\balso\b', 'ALSO'),
        (r'\bplus\b', 'PLUS'),
        (r',\s*', 'COMMA'),
    ]
    protected_text = standardized
    for span in keep_together:
        if len(span.split()) > 1:
            protected_span = span.lower().replace(' ', '_')
            protected_text = protected_text.replace(span.lower(), protected_span)
    for pattern, replacement in conjunctions:
        protected_text = re.sub(pattern, f' {replacement} ', protected_text)
    parts = re.split(r'\b(AND|TELLME|WHATABOUT|ALSO|PLUS|COMMA)\b', protected_text)
    questions = []
    current_question = []
    for part in parts:
        part = part.strip()
        if part in ['AND', 'TELLME', 'WHATABOUT', 'ALSO', 'PLUS', 'COMMA']:
            if current_question:
                questions.append(' '.join(current_question))
                current_question = []
        elif part:
            part = part.replace('_', ' ')
            current_question.append(part)
    if current_question:
        questions.append(' '.join(current_question))
    final_questions = []
    for q in questions:
        q = q.strip()
        if not q:
            continue
        if ' and ' in q and any(tok[0].isupper() for tok in q.split()):
            final_questions.append(q)
        elif not q.endswith('?'):
            if q.startswith(('what', 'how', 'who', 'why', 'when', 'where')):
                q += '?'
            else:
                q = f"What is {q}?" if len(q.split()) <= 3 else f"Tell me about {q}."
        final_questions.append(q)
    if len(final_questions) > 1:
        combined = []
        i = 0
        while i < len(final_questions):
            q = final_questions[i]
            if i < len(final_questions)-1 and ' and ' in q:
                next_q = final_questions[i+1]
                if (any(w[0].isupper() for w in q.split()) and (any(w[0].isupper() for w in next_q.split()))):
                    combined.append(f"{q} {next_q}")
                    i += 2
                    continue
            combined.append(q)
            i += 1
        final_questions = combined
    return final_questions

def extract_greeting(text):
    """Extract greeting from text if present"""
    greeting_patterns = r'\b(hello|hi|hey|good\s*(morning|afternoon|evening)|greetings|good\s*(day|night)|hey\s*there|hi\s*there|howdy|welcome)\b'
    match = re.search(greeting_patterns, text.lower(), re.IGNORECASE)
    if match:
        return match.group(0).capitalize()
    return None

def multiple_csv(files: list) -> tuple[list, dict]:
    """Read and preprocess multiple CSV files with better error handling"""
    print('Reading and preprocessing data from multiple CSV inputs...')
    chunks = []
    file_source_map = {}
    for f in files:
        try:
            print(f"Processing file: {f}")
            file_name = os.path.basename(f)
            encodings_to_try = ['utf-8', 'latin1', 'iso-8859-1', 'cp1252']
            df = None
            for encoding in encodings_to_try:
                try:
                    df = pd.read_csv(f, on_bad_lines='skip', encoding=encoding).drop_duplicates()[:1000]
                    print(f"Successfully read {f} with encoding {encoding}")
                    break
                except UnicodeDecodeError:
                    print(f"Encoding {encoding} failed for {f}")
                except Exception as read_err:
                    print(f"Error reading {f} with encoding {encoding}: {str(read_err)}")
                    if not isinstance(read_err, UnicodeDecodeError):
                        break
            if df is None:
                print(f"Could not read file {f} with any attempted encoding. Skipping.")
                continue
            df.dropna(axis=1, how='all', inplace=True)
            for _, row in df.iterrows():
                values = [f"{str(col).strip()}: {str(row[col]).strip()}"
                          for col in df.columns
                          if pd.notna(row[col]) and str(row[col]).strip()]
                chunk_text = ". ".join(values).lower()
                chunk_text = chunk_text.replace('\n', ' ').replace('\r', ' ').strip()
                if len(chunk_text) >= 20:
                    chunks.append(chunk_text)
                    file_source_map[chunk_text] = file_name
        except FileNotFoundError:
            print(f"Error: File not found {f}")
        except pd.errors.EmptyDataError:
            print(f"Error: File {f} is empty.")
        except Exception as e:
            print(f"Error processing {f}: {str(e)}")
    return chunks, file_source_map

def create_embedding_model():
    """Initialize the sentence embedding model"""
    print('Loading Embedder model...')
    model_path = os.path.abspath(os.path.join(path, 'models', 'all-MiniLM-L6-v2'))
    if not os.path.exists(model_path):
        print(f"Model path {model_path} does not exist. Please ensure the model is downloaded.")
        try:
            print("Attempting to download/load model directly from sentence_transformers...")
            return SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            print(f"Failed to load model directly: {e}")
            raise RuntimeError("Sentence embedding model could not be loaded.")
    return SentenceTransformer(model_path)

# Keep original create_faiss_index (unchanged)
def create_faiss_index(embedder, file_names):
    """Create FAISS indexes for all files, ensuring new files are embedded immediately."""
    global _EMBEDDING_UPDATE_IN_PROGRESS
    print("Creating FAISS indexes for all files...")
    all_chunks_global = []
    chunk_to_file_map_global = {}
    embeddings_dir = os.path.join(path, "embeddings")
    os.makedirs(embeddings_dir, exist_ok=True)
    embedding_dim = embedder.get_sentence_embedding_dimension()
    main_index = faiss.IndexFlatIP(embedding_dim)

    with _EMBEDDING_UPDATE_LOCK:
        if _EMBEDDING_UPDATE_IN_PROGRESS:
             print("Embedding update already in progress. Waiting...")
        _EMBEDDING_UPDATE_IN_PROGRESS = True
        print("Setting embedding update flag to True.")

    try:
        for file_path in file_names:
            if not os.path.exists(file_path):
                print(f"Warning: File {file_path} not found during FAISS index creation. Skipping.")
                continue
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            emb_path = os.path.join(embeddings_dir, f"{base_name}_faiss.index")
            chunk_path = os.path.join(embeddings_dir, f"{base_name}_chunks.npy")
            current_file_chunks = []
            current_file_embeddings = None

            if os.path.exists(emb_path) and os.path.exists(chunk_path):
                print(f"Loading existing index and chunks for {file_path}")
                try:
                    individual_index = faiss.read_index(emb_path)
                    current_file_chunks_loaded = np.load(chunk_path, allow_pickle=True).tolist()
                    if individual_index.ntotal > 0 and len(current_file_chunks_loaded) == individual_index.ntotal:
                        current_file_embeddings = np.vstack([individual_index.reconstruct(i) for i in range(individual_index.ntotal)]).astype(np.float32)
                        current_file_chunks = current_file_chunks_loaded
                    else:
                        print(f"Index/chunk mismatch or empty index for {file_path}. Rebuilding.")
                        if os.path.exists(emb_path): os.remove(emb_path)
                        if os.path.exists(chunk_path): os.remove(chunk_path)
                        current_file_embeddings = None
                except Exception as e:
                    print(f"Error loading existing index for {file_path}: {e}. Rebuilding.")
                    if os.path.exists(emb_path): os.remove(emb_path)
                    if os.path.exists(chunk_path): os.remove(chunk_path)
                    current_file_embeddings = None

            if current_file_embeddings is None:
                print(f"Creating new embeddings for {file_path} (might be new or corrupted)...")
                chunks_from_file, _ = multiple_csv([os.path.abspath(file_path)])
                if chunks_from_file:
                    current_file_chunks = chunks_from_file
                    print(f"Encoding {len(current_file_chunks)} chunks for {file_path}...")
                    embeddings_np = embedder.encode(current_file_chunks, show_progress_bar=True)
                    current_file_embeddings = normalize(embeddings_np.astype(np.float32), axis=1, norm='l2')
                    dim = current_file_embeddings.shape[1]
                    individual_index_new = faiss.IndexFlatIP(dim)
                    individual_index_new.add(current_file_embeddings)
                    faiss.write_index(individual_index_new, emb_path)
                    np.save(chunk_path, np.array(current_file_chunks, dtype=object))
                    print(f"Saved new embeddings for {file_path}")
                else:
                    print(f"No valid chunks found in {file_path}. Skipping this file for main index.")
                    continue

            if current_file_chunks and current_file_embeddings is not None and current_file_embeddings.shape[0] > 0:
                all_chunks_global.extend(current_file_chunks)
                for chunk_text_item in current_file_chunks:
                    chunk_to_file_map_global[chunk_text_item] = os.path.basename(file_path)
                main_index.add(current_file_embeddings)
            else:
                print(f"No embeddings or chunks processed for {file_path} to add to main index.")

        if main_index.ntotal == 0:
            print("Warning: Main FAISS index is empty after processing all files.")
        else:
            print(f"Main FAISS index created/updated with {main_index.ntotal} total vectors from {len([f for f in file_names if os.path.exists(f)])} processed files.")
        return embedder, all_chunks_global, main_index, chunk_to_file_map_global

    finally:
        with _EMBEDDING_UPDATE_LOCK:
            _EMBEDDING_UPDATE_IN_PROGRESS = False
            print("Embedding update completed. Flag reset to False.")

def get_context_per_question(questions, file_names, threshold=0.3, top_k=5):
    """Get context for each question separately"""
    embedder_model = create_embedding_model()
    _, all_corpus_chunks, main_faiss_index, chunk_source_map = create_faiss_index(embedder_model, file_names)
    question_contexts = {}
    all_sources = {}
    for question in questions:
        q_embed = normalize(embedder_model.encode([normalize_query(question)]), axis=1, norm='l2')
        search_k = min(top_k * 3, main_faiss_index.ntotal)
        distances, indices = main_faiss_index.search(q_embed, search_k)
        relevant_chunks = []
        for i, idx in enumerate(indices[0]):
            if 0 <= idx < len(all_corpus_chunks):
                chunk = all_corpus_chunks[idx]
                score = distances[0][i]
                source = chunk_source_map.get(chunk, "unknown_source.csv")
                if score >= threshold:
                    relevant_chunks.append((chunk, score, source))
                    all_sources[chunk] = source
        relevant_chunks.sort(key=lambda x: x[1], reverse=True)
        question_contexts[question] = "\n".join([chunk[0] for chunk in relevant_chunks[:top_k]]) 
    return question_contexts, all_sources

def history(title_id):  
    if title_id is None:  
        return []
    url = f'http://216.10.251.154:2500/chatbot/getAllChats?title_id={title_id}'
    res = requests.get(url) 
    history = res.json() 
    chat_history = []
    for entry in history['data']:
        role = "User" if entry['sender'] == 'user' else "Assistant"
        content = entry['message']
        chat_history.append({"role": role, "content": content})
    return chat_history[-4:]
    
def build_prompt(context, question, sources=None, history_text=None, is_elaboration=False):
    """Build a structured prompt that helps the model answer based strictly on context and history."""
    context = context.strip() if context else ""
    history_text = history_text.strip() if history_text else ""
    history_section = f"Previous conversation:\n{history_text}\n\n" if history_text else "No previous conversation.\n\n"
    context_section = f"Relevant information:\n{context}\n\n" if context else "No relevant information found.\n\n"
    if is_elaboration:
        instruction = f"Based on the conversation history and any relevant information provided, please elaborate further on the topic discussed or the last assistant's response. If the history is empty or irrelevant, respond that you need more context."
        question_section = f"Elaboration request: '{question}'\n\n"
    else:
        instruction = f"Based on the following relevant information and conversation history, answer the question concisely and directly. If the information does not contain the answer, state 'I don't know' and do not invent details."
        question_section = f"Question: '{question}'\n\n"
    temp = f"""You are an expert assistant. {instruction}

{history_section}{context_section}{question_section}Answer:"""
    return temp

import torch
from transformers import AutoTokenizer, AutoModelForCausalLM

def initiate_liquid_model():
    """Initialize the LiquidAI LFM2-VL-450M model"""
    print("Loading LFM2-300M model...")
    model_id = os.path.abspath(os.path.join(path , 'models',"LFM2-300M"))
    tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    dtype = torch.float16 if device == "cuda" else torch.float32
    model = AutoModelForCausalLM.from_pretrained(
        model_id,
        device_map=None,
        trust_remote_code=True,
        dtype=dtype
    ).to(device)
    model.eval()
    return tokenizer, model

_tokenizer = None
_model = None

def get_liquid_instance():
    global _tokenizer, _model
    if _model is None or _tokenizer is None:
        _tokenizer, _model = initiate_liquid_model()
    return _tokenizer, _model

def generate_liquid_response(prompt: str):
    """Generate response using LiquidAI LFM2-VL-450M model and attempt to extract the answer."""
    tokenizer, model = get_liquid_instance()
    try:
        device = next(model.parameters()).device
        inputs = tokenizer(prompt, return_tensors="pt").to(device)
        prompt_length = inputs["input_ids"].shape[1]
        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=300,
                temperature=0.1,
                top_k=40,
                top_p=0.9,
                do_sample=True,
                pad_token_id=tokenizer.eos_token_id
            )
        full_response = tokenizer.decode(outputs[0], skip_special_tokens=True)
        print(f"Full LLM Output:\n{full_response}\n--- End Full Output ---")
        answer_start_marker = "Answer:"
        prompt_end_pos = full_response.find(answer_start_marker)
        if prompt_end_pos != -1:
            extracted_answer = full_response[prompt_end_pos + len(answer_start_marker):].strip()
            extracted_answer = extracted_answer.split("Question:")[0].split("Context:")[0].split("Relevant information:")[0].split("Previous conversation:")[0].strip()
            print(f"Extracted Answer:\n{extracted_answer}\n--- End Extracted Answer ---")
            return extracted_answer
        else:
            print("Warning: 'Answer:' marker not found in LLM output. Returning full output.")
            return full_response
    except Exception as e:
        print(f"Error during LiquidAI generation: {e}")
        import traceback
        traceback.print_exc()
        return "I encountered an issue while generating a response. Please try again."

def generate_greeting_response(greeting_text):
    """Generate a friendly greeting response"""
    greeting_text_lower = greeting_text.lower()
    responses_map = {
    "hello": "Hello! How i can Help you?",
    "hi": "Hi there! How i can help you?",
    "hey": "Hey! I'm here",
    "greetings": "Greetings! How may I be of service?",
    "good morning": "Good morning! How can I help you today?",
    "good afternoon": "Good afternoon! What would you like to know?",
    "good evening": "Good evening! How may I assist you?",
    "good day": "Good day to you! How can I help?",
    "what's up": "Not much, just ready to help you! What do you need?",
    "how are you": "I'm doing well, thank you! How can I help you today?",
    "morning": "Good morning! How can I assist you today?",
    "afternoon": "Good afternoon! How may I help you?",
    "evening": "Good evening! What can I do for you?",
    "good night": "Good night! Is there anything I can help you with before you go?",
    "howdy": "Howdy! What can I help you with today?",
    "yo": "Hello! How may I assist you?",
    "hi there": "Hi there! How can I help you?",
    "hello there": "Hello there! What would you like to know?",
    "good to see you": "Good to see you too! How can I assist you today?",
    "nice to see you": "Nice to see you as well! What can I help you with?",
    "long time no see": "Hello! It's good to connect with you again. How may I help?",
    "what's new": "Nothing much here! How can I assist you today?",
    "how's it going": "Things are going well! How may I help you?",
    "how have you been": "I've been doing great, thank you for asking! How can I assist you?",
    "hope you're well": "Thank you for asking! I'm here and ready to help. What do you need?",
    "good to meet you": "Good to meet you too! How may I be of service?",
    "pleased to meet you": "Pleased to meet you as well! What can I help you with today?",
    "welcome": "Thank you! How can I assist you?",
    "hey there": "Hey there! I'm ready to help. What do you need?",
    "what's happening": "Not much on my end! How can I help you today?",
    "namaste": "Namaste! How may I be of service?",
    }
    for key_greeting, response_message in responses_map.items():
        if key_greeting in greeting_text_lower:
            return response_message
    return "Hello! I'm here to help answer your questions" 

def normalize_query(text):
    text = re.sub(r'[^\w\s]', '', text.lower())
    text = re.sub(r'\s+', ' ', text).strip()
    doc = nlp(text)
    text = ' '.join([token.lemma_ for token in doc])
    return text

def is_elaboration_requested(question):
    elaboration_patterns = [  
        r'\b(deta[il]{1,3}|detal|diteil|detials)s?\b',
        r'\b(more|morre|mor)\s*(detail[s]?|info|information|infos?)?\b',
        r'\b(additional|additinal|addtnl|extra)\s*(detail[s]?|info|information)?\b',
        r'\b(further|futher|farther)\s*(detail[s]?|info|explanation|elaboration)?\b',
        r'\bexpla[nm]{1,3}(in|ation)?\b',
        r'\b(desc[r]?ibe|descr|discribe|elaborate|elab)\s*(more|further|it|that)?\b',
        r'\b(could you|can you)?\s*(explain|describe|elaborate|expand)\s+(more|further|on that|it)?\b',
        r'\b(go on|continue|proceed|keep going)\b',
        r'\b(tell me more|say more|give me more)\b',
        r'\b(what else\??|anything else\??|more to it\??)\b',
        r'\b(expand on that|expand it|dig deeper|go deeper)\b',
        r'\b(why is that\??|how so\??|really\??)\b',
        r'\b(can you be|be)\s+(more\s+)?(specific|detailed|clear|elaborate)\b',
        r'^\s*(more|details?|info|elaborate|explain|describe|expand|further)\s*',
        r'^\s*(huh\??|hmm|hmmm|ok\??|okay\??|oh\??)\s*,?\s*(and\??|then\??|but\??)?\s*(more|details?|explain)\s*',
        r'\b(also|besides that|apart from that)\s*,?\s*(what|anything|tell me)\b',
        r'\b(moer|detailes|infomation|explane|descibe|elabrate|furthermore)\b',
        r'\b(mor details|moree info|additonal details)\b',
    ]
    question_lower = question.strip().lower()
    if len(question_lower) < 2:
        return False
    for pattern in elaboration_patterns:
        if re.search(pattern, question, flags=re.IGNORECASE):
            return True
    return False

def answer_questions(input_query, files_for_indexing=None, similarity_threshold=0.3, num_top_k=3, title_id=None):
    """Answer questions based on context from local documents."""
    if files_for_indexing is None:
        files_for_indexing = get_active_file_list()
    try:
        input_text = input_query.strip()
        if not input_text:
            return "Please ask a question.", {}, ""
        greeting = extract_greeting(input_text)
        if greeting:
            if len(input_text.strip()) == len(greeting):
                return generate_greeting_response(greeting), {}, ""
            input_text = re.sub(r'^\s*' + re.escape(greeting) + r'\s*[,.!?]?\s*', '', input_text, flags=re.IGNORECASE).strip()
        questions = extract_questions(input_text)
        contexts = {}
        sources = {}
        final_answer = ""
        source_summary = ""
        if not questions and input_text:
            questions = [input_text]
        if questions:
            if not files_for_indexing:
                print("Warning: No CSV files found in 'documents' folder. Cannot retrieve context.")
                history_list = history(title_id)
                formatted_history = ''
                for item in history_list:
                    role = "User" if item['role'] == 'User' else "Assistant"
                    content = item['content']
                    content = re.sub(r'\*{1,2}', '', content)
                    content = re.sub(r'_{1,2}', '', content)
                    formatted_history += f"{role}: {content}\n"
                answers = []
                for question in questions:
                    is_elab = is_elaboration_requested(question)
                    if is_elab and len(formatted_history.strip()) >= 5:
                        prompt = build_prompt(
                            context="",
                            question=question,
                            sources=sources,
                            history_text=formatted_history,
                            is_elaboration=True
                        )
                        if generate_liquid_response is None:
                            print("Error: generate_liquid_response is not initialized!")
                            answer = "System error: model not ready."
                        else:
                            answer = generate_liquid_response(prompt)
                    else:
                        answer = random.choice(l)
                    answers.append(answer)
            else:
                contexts, sources = get_context_per_question(questions, files_for_indexing, similarity_threshold, num_top_k)
                if not isinstance(contexts, dict) or not isinstance(sources, dict):
                    print("Context or sources not in expected format!")
                answers = []
                for i, question in enumerate(questions):
                    context = contexts.get(question, "").strip()
                    history_list = history(title_id)
                    formatted_history = ''
                    for item in history_list:
                        role = "User" if item['role'] == 'User' else "Assistant"
                        content = item['content']
                        content = re.sub(r'\*{1,2}', '', content)
                        content = re.sub(r'_{1,2}', '', content)
                        formatted_history += f"{role}: {content}\n"
                    is_elab = is_elaboration_requested(question)
                    if not context and len(formatted_history.strip()) < 5:
                        answer = random.choice(l)
                    elif not context and len(formatted_history.strip()) >= 5:
                        if is_elab:
                            prompt = build_prompt(
                                context="",
                                question=question,
                                sources=sources,
                                history_text=formatted_history,
                                is_elaboration=True
                            )
                            if generate_liquid_response is None:
                                print("Error: generate_liquid_response is not initialized!")
                                answer = "System error: model not ready."
                            else:
                                answer = generate_liquid_response(prompt)
                        else:
                            answer = random.choice(l)
                    else:
                        specific_context = context
                        prompt = build_prompt(
                            specific_context,
                            question,
                            sources,
                            formatted_history,
                            is_elaboration=is_elab
                        )
                        if generate_liquid_response is None:
                            print("Error: generate_liquid_response is not initialized!")
                            answer = "System error: model not ready."
                        else:
                            answer = generate_liquid_response(prompt)
                    if not answer:
                        answer = random.choice(l)
                    answers.append(answer)
            if len(answers) == 1:
                final_answer = answers[0]
            else:
                final_answer = " ".join(f"• {ans}" for ans in answers)
            if greeting:
                greeting_response = generate_greeting_response(greeting)
                final_answer = f"{greeting_response}\n{final_answer}"
            if files_for_indexing:
                unique_sources = set(s for s in sources.values() if s and s.lower() != "unknown_source.csv")
                if unique_sources:
                    source_summary = "Information based on: " + ", ".join(sorted(unique_sources))
            else:
                source_summary = "No documents available for context."
        else:
            final_answer = "I couldn't identify a specific question. Please try rephrasing."
        return final_answer, contexts, source_summary
    except Exception as e:
        print(f"Error in answer_questions: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"An error occurred while processing your questions: {str(e)}", {}, ""

_external_docs_cache = {}
_external_docs_last_fetch = 0
EXTERNAL_DOCS_CACHE_DURATION = 300
_data_files_cache_list = None
_data_files_last_update_time = 0

def get_active_file_list():
    """Get file names from BOTH local 'documents' folder and external API"""
    global _data_files_cache_list, _data_files_last_update_time
    cache_validity_duration = 300
    if _data_files_cache_list is not None and (time.time() - _data_files_last_update_time) < cache_validity_duration:
        print("Using cached file list.")
        return _data_files_cache_list
    print("Building active file list from local and external sources...")
    documents_dir = os.path.join(path, 'documents')
    active_files = []
    try:
        print("\n--- Syncing External Documents ---")
        external_csv_files = sync_external_documents()
        print(f"External sync complete: {len(external_csv_files)} files")
    except Exception as e:
        print(f"Error syncing external documents: {e}")
        import traceback
        traceback.print_exc()
        external_csv_files = []
    if os.path.exists(documents_dir):
        print("\n--- Scanning Local Documents Folder ---")
        for filename in os.listdir(documents_dir):
            if filename.lower().endswith('.csv') and not filename.startswith('.'):
                full_path = os.path.join(documents_dir, filename)
                if os.path.isfile(full_path) and full_path not in active_files:
                    active_files.append(full_path)
        print(f"Local scan complete: {len(active_files)} total CSV files found")
    else:
        print(f"Documents directory '{documents_dir}' does not exist.")
    _data_files_cache_list = active_files
    _data_files_last_update_time = time.time()
    print(f"\n✓ Total active files: {len(active_files)}")
    return active_files

def fetch_external_documents_list(api_url="http://15.206.70.213:7501/documents/getAllDocuments"):
    """Fetch list of available documents from external API"""
    global _external_docs_cache, _external_docs_last_fetch
    current_time = time.time()
    if _external_docs_cache and (current_time - _external_docs_last_fetch) < EXTERNAL_DOCS_CACHE_DURATION:
        print("Using cached external documents list.")
        return _external_docs_cache.get('documents', [])
    try:
        print(f"Fetching documents list from: {api_url}")
        response = requests.get(api_url, timeout=10)
        response.raise_for_status()
        data = response.json()
        if data.get('success') == 'success' and 'data' in data:
            documents = data['data']
            _external_docs_cache['documents'] = documents
            _external_docs_last_fetch = current_time
            print(f"Successfully fetched {len(documents)} documents from external API.")
            return documents
        else:
            print(f"Unexpected response format from external API: {data}")
            return []
    except requests.exceptions.RequestException as e:
        print(f"Error fetching external documents list: {e}")
        return []
    except Exception as e:
        print(f"Unexpected error fetching documents list: {e}")
        return []

def download_external_document(doc_name, base_url="http://15.206.70.213:7501/documents/"):
    """Download a document from external API"""
    try:
        encoded_name = quote(doc_name)
        doc_url = f"{base_url}{encoded_name}"
        print(f"Downloading document from: {doc_url}")
        response = requests.get(doc_url, timeout=30)
        response.raise_for_status()
        print(f"Successfully downloaded: {doc_name} ({len(response.content)} bytes)")
        return response.content
    except requests.exceptions.RequestException as e:
        print(f"Error downloading document {doc_name}: {e}")
        return None
    except Exception as e:
        print(f"Unexpected error downloading {doc_name}: {e}")
        return None

def convert_docx_to_csv(docx_content, output_csv_path):
    """Convert DOCX content to CSV format"""
    try:
        doc = DocxDocument(io.BytesIO(docx_content))
        rows = []
        row_id = 1
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 10:
                rows.append({
                    'id': row_id,
                    'type': 'paragraph',
                    'content': text
                })
                row_id += 1
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cell_texts = [cell.text.strip() for cell in row.cells]
                combined_text = " | ".join([txt for txt in cell_texts if txt])
                if combined_text:
                    rows.append({
                        'id': row_id,
                        'type': f'table_{table_idx + 1}',
                        'content': combined_text
                    })
                    row_id += 1
        if rows:
            df = pd.DataFrame(rows)
            df.to_csv(output_csv_path, index=False, encoding='utf-8')
            print(f"Converted DOCX to CSV: {output_csv_path} ({len(rows)} rows)")
            return True
        else:
            print(f"No content extracted from DOCX")
            return False
    except Exception as e:
        print(f"Error converting DOCX to CSV: {e}")
        import traceback
        traceback.print_exc()
        return False

def get_file_hash(file_path):
    """Calculate MD5 hash of a file to check if it has changed"""
    try:
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        print(f"Error calculating hash for {file_path}: {e}")
        return None

def sync_external_documents():
    """
    MODIFIED: Syncs documents and queues new ones for background embedding
    """
    documents_dir = os.path.join(path, 'documents')
    os.makedirs(documents_dir, exist_ok=True)
    metadata_file = os.path.join(documents_dir, '.external_docs_metadata.json')
    existing_metadata = {}
    if os.path.exists(metadata_file):
        try:
            with open(metadata_file, 'r') as f:
                existing_metadata = json.load(f)
        except Exception as e:
            print(f"Error loading metadata file: {e}")
    external_docs = fetch_external_documents_list()
    if not external_docs:
        print("No external documents found or API unavailable.")
        return []
    synced_csv_files = []
    updated_metadata = {}
    new_files = []
    for doc_info in external_docs:
        doc_name = doc_info.get('name')
        doc_id = doc_info.get('id')
        if not doc_name:
            continue
        print(f"\n{'='*60}")
        print(f"Processing: {doc_name} (ID: {doc_id})")
        print(f"{'='*60}")
        base_name = os.path.splitext(doc_name)[0]
        csv_filename = f"{base_name}.csv"
        csv_path = os.path.join(documents_dir, csv_filename)
        file_exists = os.path.exists(csv_path)
        skip_download = False
        if file_exists:
            if str(doc_id) in existing_metadata:
                old_hash = existing_metadata[str(doc_id)].get('hash')
                current_hash = get_file_hash(csv_path)
                if old_hash == current_hash:
                    print(f"✓ File already exists and unchanged: {csv_filename}")
                    skip_download = True
                    synced_csv_files.append(csv_path)
                    updated_metadata[str(doc_id)] = existing_metadata[str(doc_id)]
        if not skip_download:
            doc_content = download_external_document(doc_name)
            if doc_content is None:
                print(f"✗ Failed to download: {doc_name}")
                continue
            file_ext = os.path.splitext(doc_name)[1].lower()
            if file_ext == '.csv':
                with open(csv_path, 'wb') as f:
                    f.write(doc_content)
                print(f"✓ Saved CSV directly: {csv_filename}")
            elif file_ext in ['.docx', '.doc']:
                success = convert_docx_to_csv(doc_content, csv_path)
                if not success:
                    print(f"✗ Failed to convert DOCX to CSV: {doc_name}")
                    continue
                print(f"✓ Converted and saved: {csv_filename}")
            else:
                print(f"✗ Unsupported file format: {file_ext}")
                continue
            file_hash = get_file_hash(csv_path)
            updated_metadata[str(doc_id)] = {
                'name': doc_name,
                'csv_path': csv_path,
                'hash': file_hash,
                'last_synced': time.time()
            }
            synced_csv_files.append(csv_path)
            new_files.append(csv_path)
    try:
        with open(metadata_file, 'w') as f:
            json.dump(updated_metadata, f, indent=2)
        print(f"\n✓ Metadata saved: {metadata_file}")
    except Exception as e:
        print(f"Error saving metadata: {e}")
    
    # NEW: Queue new files for background embedding
    for file_path in new_files:
        queue_file_for_embedding(file_path)
    
    print(f"\n{'='*60}")
    print(f"Sync Complete: {len(synced_csv_files)} documents ready")
    if new_files:
        print(f"New files queued for embedding: {len(new_files)}")
    print(f"{'='*60}\n")
    return synced_csv_files

def response_handler(response):
    sentences = re.split(r'\n+|(?<=\d)\. |\.\s+', response)
    unique_sentences = []
    seen_sentences = set()
    for sentence in sentences:
        normalized_sentence = sentence.strip().lower()
        if normalized_sentence and normalized_sentence not in seen_sentences:
            unique_sentences.append(sentence.strip())
            seen_sentences.add(normalized_sentence)
    cleaned_response = ". ".join(unique_sentences)
    if cleaned_response and not cleaned_response.endswith('.'):
        cleaned_response += "."
    return cleaned_response

def extract_keywords(text):
    stopwords = set(["the", "is", "and", "of", "in", "on", "for", "a", "an", "to", "by", "with", "under", "that", "as", "has", "are", "this", "which", "be"])
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    keywords = set([word for word in words if word not in stopwords])
    return keywords

def apply_formatting(text, keywords):
    """Format the keywords in the text using bold and italics."""
    def format_word(word):
        return f"**{word.capitalize()}**" if len(word) > 5 else f"*{word}*"
    for kw in sorted(keywords, key=lambda x: -len(x)):
        pattern = re.compile(rf'\b({kw})\b', re.IGNORECASE)
        text = pattern.sub(lambda m: format_word(m.group(1)), text)
    return text

def text_formatting(llm_answer, question, contexts):
    if isinstance(contexts, dict):
        all_context_text = " ".join(contexts.values())
    else:
        all_context_text = str(contexts)
    question_keywords = extract_keywords(question)
    context_keywords = extract_keywords(all_context_text)
    all_keywords = question_keywords.union(context_keywords)
    formatted_text = apply_formatting(llm_answer, all_keywords)
    return formatted_text

class FileList(BaseModel):
    file_names: list[str]

from typing import List
class FilePaths(BaseModel):
    paths: List[str]

@app.post("/remove_by_paths")
def remove_files_by_paths(file_paths: FilePaths):  
    print("**************** file paths *****************") 
    print(file_paths)
    print('*********************************************')
    doc_dir = os.path.join(path, 'documents')
    embed_dir = os.path.join(path, 'embeddings') 
    print("*****************************") 
    print(doc_dir , embed_dir) 
    print('*****************************')
    removed = []
    not_found = []
    for raw_path in file_paths.paths:
        file_name = os.path.basename(raw_path.replace("\\", "/"))
        base_name = os.path.splitext(file_name)[0]
        doc_path = os.path.join(doc_dir, file_name)
        index_path = os.path.join(embed_dir, f'{base_name}_faiss.index')
        chunks_path = os.path.join(embed_dir, f'{base_name}_chunks.npy')
        for file_path_to_remove in [doc_path, index_path, chunks_path]:
            try:
                if os.path.exists(file_path_to_remove):
                    os.remove(file_path_to_remove)
                    removed.append(file_path_to_remove)
                else:
                    not_found.append(file_path_to_remove)
            except Exception as e:
                print(f"Error deleting {file_path_to_remove}: {e}")
                not_found.append(file_path_to_remove)
    return {
        "deleted_files": removed,
        "not_found": not_found
    }

spell = SpellChecker()
custom_words = {"Thermax" ,'Boilers' , 'Efficiency', 'Marworx' , 'Refractory','Castable' , 'Tastable'} 
spell.word_frequency.load_words(custom_words) 
    
def correct_spelling_custom_dict(text, spell_checker):
    words = text.split()
    corrected_words = []
    for word in words:
        if word.lower() in spell_checker:
            corrected_words.append(word) 
        else:
            correction = spell_checker.correction(word.lower())
            if correction:
                if word.isupper():
                    corrected_words.append(correction.upper())
                elif word.istitle():
                     corrected_words.append(correction.capitalize())
                else:
                    corrected_words.append(correction)
            else:
                 corrected_words.append(word)
    return " ".join(corrected_words)

class QueryRequest(BaseModel):
    query: Optional[str] = None

# ==================== FASTAPI STARTUP/SHUTDOWN ====================
@app.on_event("startup")
async def startup_event():
    """Start background worker when server starts"""
    start_background_worker()
    print("✅ Server started with background embedding worker")

@app.on_event("shutdown")
async def shutdown_event():
    """Stop background worker when server shuts down"""
    stop_background_worker()
    print("✅ Server shutdown complete")

# ==================== NEW ENDPOINTS ====================
@app.get("/embedding_status")
async def get_embedding_status():
    """Check the status of background embedding generation"""
    status = embedding_status.get_status()
    return JSONResponse({
        "status": "running" if _BACKGROUND_THREAD and _BACKGROUND_THREAD.is_alive() else "stopped",
        "queue_size": _EMBEDDING_QUEUE.qsize(),
        "pending_files": len(status["pending"]),
        "processing": status["processing"],
        "completed": len(status["completed"]),
        "failed": len(status["failed"]),
        "details": status
    })

@app.post("/force_embedding")
async def force_embedding_generation(file_paths: List[str]):
    """Force immediate embedding generation for specific files"""
    try:
        queued = []
        failed = []
        for file_path in file_paths:
            full_path = os.path.join(path, "documents", file_path)
            if queue_file_for_embedding(full_path):
                queued.append(file_path)
            else:
                failed.append(file_path)
        return JSONResponse({
            "success": True,
            "queued": queued,
            "failed": failed,
            "message": f"Queued {len(queued)} files for embedding generation"
        })
    except Exception as e:
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

# ==================== MODIFIED MAIN ENDPOINT ====================
@app.api_route("/get_info", methods=["GET", "POST"])
async def ask(request: Request, query: Optional[str] = Form(None)):
    """MODIFIED: Now uses fast index loading with background embedding"""
    global _EMBEDDING_UPDATE_IN_PROGRESS
    try:
        user_query = None
        title_id = None
        if request.method == "GET":
            user_query = request.query_params.get("query")
            title_id = 115
            if title_id is not None:
                try:
                    title_id = int(title_id)
                except ValueError:
                    title_id = None
        elif request.method == "POST":
            if request.headers.get("content-type", "").startswith("application/json"):
                json_data = await request.json()
                user_query = json_data.get("query")
                title_id = json_data.get("title_id")
            else:
                form = await request.form()
                user_query = form.get("query")
                title_id = form.get("title_id")
            if title_id is not None:
                try:
                    title_id = int(title_id)
                except ValueError:
                    title_id = None
        if not user_query or not user_query.strip():
            return JSONResponse(
                {"error": "Missing or empty question parameter 'query'"},
                status_code=400
            )
        is_follow_up = is_elaboration_requested(user_query.strip())
        if is_follow_up and title_id is None:
            return JSONResponse({
                "question": user_query,
                "answer": "I can't provide more details because I don't know which conversation you're referring to. Please include the chat ID (title_id).",
                "context": {},
                "context_source_summary": "",
                "time_taken": 0.01
            }, status_code=400)
        request_start_time = time.time()
        
        # Use fast index loading
        current_active_files = get_active_file_list()
        if not current_active_files:
            return JSONResponse({
                "question": user_query,
                "answer": "No data documents available. Please try again later.",
                "context_source_summary": "No data files found.",
                "time_taken": round(time.time() - request_start_time, 2)
            })
        
        print("--"*30)
        print(f"Query: {user_query} | Title ID: {title_id}")
        print("--"*30)
        
        # Use FAST index loading (doesn't block)
        embedder = create_embedding_model()
        _, chunks, index, source_map = create_faiss_index_fast(embedder, current_active_files)
        
        if index.ntotal == 0:
            return JSONResponse({
                "question": user_query,
                "answer": "Documents are being processed. Please try again in a moment.",
                "context_source_summary": "Embeddings being generated in background",
                "time_taken": round(time.time() - request_start_time, 2),
                "embedding_status": "/embedding_status"
            })
        
        llm_answer, contexts, context_summary = answer_questions(
            user_query,
            files_for_indexing=current_active_files,
            similarity_threshold=0.21,
            num_top_k=5,
            title_id=title_id
        )
        llm_answer = correct_spelling_custom_dict(llm_answer, spell)
        llm_answer = response_handler(llm_answer)
        llm_answer = re.sub(r'(?<!)(\d+\.\s*)', r'\1', llm_answer)
        llm_answer = text_formatting(llm_answer, user_query, contexts)
        request_end_time = time.time()
        return {
            "question": user_query,
            "answer": llm_answer.title() if isinstance(llm_answer, str) and llm_answer else llm_answer,
            "context": contexts,
            "context_source_summary": context_summary,
            "time_taken": round(request_end_time - request_start_time, 2),
            "indexed_documents": index.ntotal > 0
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"error": f"An internal server error occurred: {str(e)}"},
            status_code=500
        )

class AddKnowledgeRequest(BaseModel):
    question: str
    answer: str
    target_file_base_name: str = None

def _find_any_existing_csv_base_name():
    """Finds the base name of any existing CSV file in the documents directory."""
    documents_dir = os.path.join(path, "documents")
    if not os.path.exists(documents_dir):
        print(f"Documents directory not found: {documents_dir}")
        return None
    try:
        for filename in os.listdir(documents_dir):
            if filename.lower().endswith('.csv'):
                base_name = os.path.splitext(filename)[0]
                emb_path = os.path.join(path, "embeddings", f"{base_name}_faiss.index")
                chunk_path = os.path.join(path, "embeddings", f"{base_name}_chunks.npy")
                if os.path.exists(emb_path) and os.path.exists(chunk_path):
                    print(f"Found existing CSV with embeddings: {base_name}")
                    return base_name
        for filename in os.listdir(documents_dir):
             if filename.lower().endswith('.csv'):
                base_name = os.path.splitext(filename)[0]
                print(f"Found existing CSV (no embeddings yet): {base_name}")
                return base_name
    except Exception as e:
        print(f"Error finding an existing CSV file: {e}")
    return None

def _add_knowledge_to_specific_index(embedder, target_base_name, question_text, answer_text):
    """Adds a new Q&A pair to the FAISS index and chunks file for a specific CSV."""
    try:
        embeddings_dir = os.path.join(path, "embeddings")
        os.makedirs(embeddings_dir, exist_ok=True)
        emb_path = os.path.join(embeddings_dir, f"{target_base_name}_faiss.index")
        chunk_path = os.path.join(embeddings_dir, f"{target_base_name}_chunks.npy")
        new_chunk_text =  f"{answer_text.strip()} This addresses the topic: {question_text.strip()}.".lower()
        new_chunk_text = new_chunk_text.replace('\n', ' ').replace('\r', ' ').strip()
        if len(new_chunk_text) < 20:
             print("Warning: New chunk is too short. Skipping addition.")
             return False, "New knowledge chunk is too short."
        index = None
        chunks_list = []
        dim = embedder.get_sentence_embedding_dimension()
        if os.path.exists(emb_path) and os.path.exists(chunk_path):
            try:
                print(f"Loading existing index and chunks for {target_base_name}...")
                index = faiss.read_index(emb_path)
                chunks_list = np.load(chunk_path, allow_pickle=True).tolist()
                print(f"Loaded index with {index.ntotal} vectors.")
                if index.d != dim:
                    raise ValueError(f"Dimension mismatch: index.d={index.d}, expected={dim}")
                if index.ntotal != len(chunks_list):
                     print(f"Warning: Index/vector count mismatch (index: {index.ntotal}, chunks: {len(chunks_list)}). Rebuilding index.")
                     index = None
                     chunks_list = []
            except Exception as e:
                print(f"Error loading existing index/chunks for {target_base_name}: {e}. Will create new ones.")
                index = None
                chunks_list = []
        if index is None:
            print(f"Creating new FAISS index for {target_base_name}...")
            index = faiss.IndexFlatIP(dim)
            chunks_list = []
        if new_chunk_text in chunks_list:
            print("Warning: This knowledge chunk already exists in the target index. Skipping addition.")
            return False, "Knowledge chunk already exists."
        print("Generating embedding for new knowledge...")
        start_time = time.time()
        new_embedding = embedder.encode([new_chunk_text])
        end_time = time.time()
        print(f"Embedding generation took {end_time - start_time:.2f} seconds.")
        new_embedding_normalized = normalize(new_embedding.astype(np.float32), axis=1, norm='l2')
        print("Adding new embedding and chunk...")
        index.add(new_embedding_normalized)
        chunks_list.append(new_chunk_text)
        print("Saving updated FAISS index and chunks...")
        faiss.write_index(index, emb_path)
        np.save(chunk_path, np.array(chunks_list, dtype=object))
        print("Knowledge added and saved successfully.")
        return True, f"Knowledge added successfully to {target_base_name}."
    except Exception as e:
        error_msg = f"Error adding knowledge to {target_base_name}: {e}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        return False, error_msg
    
@app.post("/add_qa")
async def add_knowledge_endpoint(request_data: AddKnowledgeRequest):
    """Endpoint to add new verified question-answer pairs."""
    if not request_data.question.strip() or not request_data.answer.strip():
        raise HTTPException(status_code=400, detail="Both 'question' and 'answer' are required and cannot be empty.")
    target_base_name = request_data.target_file_base_name
    selected_base_name = None
    if target_base_name:
        expected_csv_path = os.path.join(path, "documents", f"{target_base_name}.csv")
        emb_path = os.path.join(path, "embeddings", f"{target_base_name}_faiss.index")
        chunk_path = os.path.join(path, "embeddings", f"{target_base_name}_chunks.npy")
        if os.path.exists(expected_csv_path):
            selected_base_name = target_base_name
            print(f"Target file specified and found: {selected_base_name}")
        else:
            print(f"Target file '{target_base_name}' not found (expected {expected_csv_path}). Will select another file.")
    if not selected_base_name:
        print("No valid target file specified or found. Searching for any existing CSV...")
        selected_base_name = _find_any_existing_csv_base_name()
        if not selected_base_name:
            raise HTTPException(status_code=400, detail="No target file specified and no existing CSV files found to add knowledge to.")
    try:
        print("Initializing embedding model...")
        embedder = create_embedding_model()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to initialize embedding model: {e}")
    print(f"Attempting to add knowledge to '{selected_base_name}'...")
    success, message = _add_knowledge_to_specific_index(
        embedder=embedder,
        target_base_name=selected_base_name,
        question_text=request_data.question,
        answer_text=request_data.answer
    )
    if not success:
        if "already exists" in message or "too short" in message:
             raise HTTPException(status_code=400, detail=message)
        else:
             raise HTTPException(status_code=500, detail=message)
    return {"message": message, "target_file": f"{selected_base_name}.csv"}

@app.post("/sync_external_docs")
async def trigger_external_sync():
    """Manually trigger synchronization of external documents"""
    try:
        print("Manual sync triggered via API endpoint")
        global _data_files_cache_list, _data_files_last_update_time, _external_docs_cache, _external_docs_last_fetch
        _data_files_cache_list = None
        _data_files_last_update_time = 0
        _external_docs_cache = {}
        _external_docs_last_fetch = 0
        synced_files = sync_external_documents()
        all_files = get_active_file_list()
        return JSONResponse({
            "success": True,
            "message": "External documents synced successfully",
            "synced_files_count": len(synced_files),
            "total_files_count": len(all_files),
            "synced_files": [os.path.basename(f) for f in synced_files],
            "all_files": [os.path.basename(f) for f in all_files]
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": f"Sync failed: {str(e)}"},
            status_code=500
        )

@app.get("/list_documents")
async def list_all_documents():
    """List all available documents (local + external synced)"""
    try:
        documents_dir = os.path.join(path, 'documents')
        embeddings_dir = os.path.join(path, 'embeddings')
        all_files = get_active_file_list()
        document_info = []
        for file_path in all_files:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            emb_path = os.path.join(embeddings_dir, f"{base_name}_faiss.index")
            chunk_path = os.path.join(embeddings_dir, f"{base_name}_chunks.npy")
            has_embeddings = os.path.exists(emb_path) and os.path.exists(chunk_path)
            file_stats = os.stat(file_path)
            document_info.append({
                "filename": os.path.basename(file_path),
                "full_path": file_path,
                "has_embeddings": has_embeddings,
                "file_size_bytes": file_stats.st_size,
                "last_modified": file_stats.st_mtime
            })
        return JSONResponse({
            "success": True,
            "total_documents": len(document_info),
            "documents": document_info
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": f"Failed to list documents: {str(e)}"},
            status_code=500
        )

if __name__ == '__main__': 
    import uvicorn
    uvicorn.run(app, host='0.0.0.0', port=5000)